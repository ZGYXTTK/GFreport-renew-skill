# -*- coding: utf-8 -*-
"""
make_fixtures.py —— 生成 evals/cases/fixtures/ 自举冒烟集（幂等）

产出：
  sample_old.docx        旧月报样本（1 标题 + 1 段落 + 在审表 5 行 + 融资表 2 行）
  sample_new.docx        新月报样本（正文/数值更新；融资表含 1 条重复事件 + 1 条无名主体）
  source_sample.csv      源 CSV（溯源回读目标）
  溯源.jsonl             正向溯源：5 条全部可在 docx 与 CSV 双侧回读一致
  溯源_bad.jsonl         负向溯源：1 条幽灵条目 + 1 条错值（门禁必须拦截）

用途：evals/gfreport-renew.eval.md 的 C1/C2/C2b/C3/C3b/C4 冒烟；
     正样本不误杀（exit 0），负样本必拦截（exit 1）。
"""
import csv
import json
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError as e:
    raise SystemExit('❌ python-docx 未安装：pip install python-docx') from e

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

_HERE = Path(__file__).parent
_FIXTURES = _HERE / 'fixtures'
_FIXTURES.mkdir(parents=True, exist_ok=True)

_COMPANIES = [
    ('A 公司', '科创板', '已问询', '0.5 亿元'),
    ('B 公司', '创业板', '已受理', '0.8 亿元'),
    ('C 公司', '主板', '注册', '1.5 亿元'),
    ('D 公司', '科创板', '上市委', '2.0 亿元'),
    ('E 公司', '北交所', '已问询', '0.3 亿元'),
]


def _build(doc_head, para, companies, funding_rows):
    doc = Document()
    doc.add_heading(doc_head, level=1)
    doc.add_paragraph(para)
    t1 = doc.add_table(rows=1 + len(companies), cols=4)
    for j, h in enumerate(('公司简称', '板块', '在审进度', '最新融资额')):
        t1.rows[0].cells[j].text = h
    for i, row in enumerate(companies, start=1):
        for j, v in enumerate(row):
            t1.rows[i].cells[j].text = v
    t2 = doc.add_table(rows=1 + len(funding_rows), cols=4)
    for j, h in enumerate(('公司简称', '轮次', '金额', '投资方')):
        t2.rows[0].cells[j].text = h
    for i, row in enumerate(funding_rows, start=1):
        for j, v in enumerate(row):
            t2.rows[i].cells[j].text = v
    return doc


def make():
    # ---- sample_old：2026-07 期 ----
    old_funding = [
        ('经纬测试', '天使轮', '1.00 亿元', '闻名投资'),
        ('匿名样本', '天使轮', '0.30 亿元', '某产业机构'),
    ]
    _build('2026-07 月报 · 样本',
           '本期共 5 家在审企业，环比增长 12%。融资金额 1.2 亿元。',
           _COMPANIES, old_funding
           ).save(str(_FIXTURES / 'sample_old.docx'))

    # ---- sample_new：2026-08 期（C 公司金额 1.2→1.5；融资表重复 + 无名主体） ----
    new_funding = [
        ('经纬测试', '天使轮', '1.00 亿元', '闻名投资'),   # 与上期同主体同轮次 → 重复 ⚠️
        ('匿名样本', '天使轮', '0.30 亿元', '某产业机构'),  # 无名主体 → ❌
    ]
    _build('2026-08 月报 · 样本',
           '本期共 5 家在审企业，环比持平。融资金额 1.5 亿元。',
           _COMPANIES, new_funding
           ).save(str(_FIXTURES / 'sample_new.docx'))

    # ---- source_sample.csv ----
    with open(_FIXTURES / 'source_sample.csv', 'w', encoding='utf-8', newline='') as f:
        w = csv.writer(f)
        w.writerow(['company', 'progress', 'amount_yi'])
        for name, _b, prog, amt in _COMPANIES:
            w.writerow([name, prog, amt.replace(' 亿元', '')])

    # ---- 溯源.jsonl（正向：全部可双侧回读一致） ----
    recs = []
    for name, _b, prog, amt in _COMPANIES:
        recs.append({
            'section': '在审企业', 'metric': f'{name} 在审进度', 'value': prog, 'unit': '',
            'cell': '在审企业', 'source_file': 'source_sample.csv', 'source_field': 'progress',
            'source_key': name, 'source_url': 'fixture', 'source_type': 'fixture',
            'as_of': '2026-08-31', 'cross_checked': True, 'cross_source': 'fixture 双源',
        })
    with open(_FIXTURES / '溯源.jsonl', 'w', encoding='utf-8') as f:
        for r in recs:
            f.write(json.dumps(r, ensure_ascii=False) + '\n')

    # ---- 溯源_bad.jsonl（负向：幽灵 + 错值，门禁必须拦截） ----
    bad = [
        {**recs[0], 'metric': '峰飞测试 V5000 在审进度', 'source_key': '峰飞测试'},  # 幽灵：docx 无
        {**recs[2], 'value': '已受理'},                                            # 错值：docx=注册
    ]
    with open(_FIXTURES / '溯源_bad.jsonl', 'w', encoding='utf-8') as f:
        for r in bad:
            f.write(json.dumps(r, ensure_ascii=False) + '\n')

    print('[OK] fixtures refreshed:', _FIXTURES)
    print('     sample_old.docx / sample_new.docx / source_sample.csv / 溯源.jsonl / 溯源_bad.jsonl')


if __name__ == '__main__':
    make()
