# -*- coding: utf-8 -*-
"""
content_freshness.py —— ⑩ 内容新鲜度门禁（v2：行业关键词 pack 化）

v0.1.0 教训：v1 新月报"格式保真 100% + 门禁全过"但内容 90% 仍是上期旧文。
本门禁：通过关键词命中数验证新月报确实包含目标月份的最新数据。

v2 修复（2026-09-01 审计）：
  v1 把航空航天 2026-08 期的 44 个关键词硬编码在通用引擎里（DEFAULT_KEYWORDS_BY_YM），
  换行业/换月份即退化为 2 条泛化关键词，形同虚设——冒烟测试中 fixtures 样本被航天词表
  误判 0 命中即是现场重现。v2 引擎不再内置任何行业词表，关键词表一律从
  packs/<pack>/config/新鲜度关键词.json 加载（结构：{"<YYYY-MM>": [...], "default": [...]}）；
  未配置时退化为 [ym, "YYYY 年 M 月"] 2 条并在报告标注「泛化模式」。

用法：
  python scripts/audit/content_freshness.py <new.docx> --ym 2026-08 --pack aerospace
  python scripts/audit/content_freshness.py <new.docx> --ym 2026-08 --old <old.docx>

退出码：
  0 = 内容已更新（命中 ≥ 阈值）
  1 = 内容未更新（命中 < 阈值）
"""
import argparse
import json
import re
import sys
from pathlib import Path

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
except ImportError:
    raise SystemExit('❌ python-docx 未安装')

_HERE = Path(__file__).parent
_BASE = _HERE.parent


def _load_keywords(pack, ym):
    """从 packs/<pack>/config/新鲜度关键词.json 加载 {ym: [...]}；退化为 ym 泛化 2 条。"""
    cfg = _BASE / 'packs' / pack / 'config' / '新鲜度关键词.json'
    generic = [ym, ym.replace('-', ' 年 ') + ' 月']
    if cfg.exists():
        try:
            data = json.loads(cfg.read_text(encoding='utf-8'))
            kws = data.get(ym) or data.get('default') or []
            if kws:
                return kws, f'packs/{pack}/config/新鲜度关键词.json'
        except (json.JSONDecodeError, OSError):
            pass
    return generic, None


def _extract_text(docx_path):
    """提取 docx 全部文本（含段落+表格）。"""
    doc = Document(str(docx_path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return parts


def _count_keywords(texts, keywords):
    """统计每个 keyword 命中次数。"""
    big = '\n'.join(texts)
    return {kw: big.count(kw) for kw in keywords}


def main():
    ap = argparse.ArgumentParser(description='内容新鲜度门禁：验证新月报包含目标月份数据')
    ap.add_argument('new_docx', help='新月报 docx 路径')
    ap.add_argument('--ym', required=True, help='目标月份 YYYY-MM')
    ap.add_argument('--pack', default='_default', help='行业包名（加载 packs/<pack>/config/新鲜度关键词.json）')
    ap.add_argument('--old', default=None, help='上期月报（用于检测是否复制旧文）')
    ap.add_argument('--threshold', type=float, default=0.5,
                    help='关键词命中率阈值（默认 50%%，每条至少命中 1 次）')
    ap.add_argument('--out', default='content_freshness.md')
    args = ap.parse_args()

    keywords, source = _load_keywords(args.pack, args.ym)
    generic_mode = source is None
    if generic_mode:
        print(f'⚠️  packs/{args.pack} 未配置 {args.ym} 关键词，使用泛化 2 条（{keywords}）；'
              f'建议在 packs/{args.pack}/config/新鲜度关键词.json 配置行业月度词表')

    texts = _extract_text(args.new_docx)
    counts = _count_keywords(texts, keywords)
    hit = sum(1 for n in counts.values() if n > 0)
    total = len(keywords)
    hit_rate = hit / total if total else 0

    lines = [
        '# 内容新鲜度报告（content_freshness.py v2）',
        f'来源：{args.new_docx} ｜ 目标月份：{args.ym} ｜ 关键词总数：{total}',
        f'词表来源：{source or "泛化模式（ym 字符串 2 条）｜pack: " + args.pack}，'
        f'命中关键词：{hit} / {total} = {hit_rate:.1%}（阈值 {args.threshold:.0%}）',
        '',
        '## 关键词命中明细',
        '| 关键词 | 命中数 |',
        '| --- | --- |',
    ]
    for kw, n in sorted(counts.items(), key=lambda x: -x[1]):
        lines.append(f'| {kw} | {n} |')

    if args.old:
        old_texts = _extract_text(args.old)
        old_counts = _count_keywords(old_texts, keywords)
        new_only = []
        old_only = []
        for kw in keywords:
            new_n = counts.get(kw, 0)
            old_n = old_counts.get(kw, 0)
            if new_n > 0 and old_n == 0:
                new_only.append(kw)
            elif new_n == 0 and old_n > 0:
                old_only.append(kw)
        lines += [
            '',
            '## 与上期月报对比（仅 8 月新词）',
            f'新月报独有 8 月词：{", " .join(new_only) if new_only else "（无）"}',
            f'新月报缺失（仅上期有）：{", " .join(old_only) if old_only else "（无）"}',
        ]

    passed = hit_rate >= args.threshold

    Path(args.out).write_text('\n'.join(lines), encoding='utf-8')

    print(f'\n📊 内容新鲜度：{hit}/{total} = {hit_rate:.1%}（阈值 {args.threshold:.0%}）')
    print(f'   报告：{args.out}')
    sys.exit(0 if passed else 1)


if __name__ == '__main__':
    main()
