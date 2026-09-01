# -*- coding: utf-8 -*-
"""
extract_numbers.py —— ① 数字提取门禁（v1）

从旧月报 docx 提取「结论性语句 + 数字」，产出清单（不阻断）。
依赖 python-docx（缺失时报错而非静默）。
"""
import argparse
import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError as e:
    raise SystemExit('❌ python-docx 未安装：pip install python-docx') from e

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

# 匹配数字（含单位）：1,234 / 12.34 / 1,234.56 / 12亿 / 35%
_NUM = re.compile(r'(\d[\d,]*(?:\.\d+)?)(\s*)([%亿万元美元港元人民币千百]*)')

# 含数字的"结论性"句子启发式：包含数量/家/笔/亿元/%/同比/环比 等关键词
_CLAIM_HINTS = re.compile(
    r'(同比|环比|共\d|家|笔|亿元|万元|亿美元|%\d|增\w|减\w|达\w|超过|突破|增长|下降)',
    re.IGNORECASE
)


def _para_text(p):
    return ''.join(r.text or '' for r in p.runs)


def extract(docx_path):
    doc = Document(str(docx_path))
    rows = []
    n = 0
    for pi, p in enumerate(doc.paragraphs):
        text = _para_text(p).strip()
        if not text:
            continue
        if not _CLAIM_HINTS.search(text):
            continue
        nums = _NUM.findall(text)
        if not nums:
            continue
        rows.append({'段落号': pi, '类型': 'paragraph', '内容': text[:200], '数字数量': len(nums)})
        n += 1

    # 表格
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if not txt or not _CLAIM_HINTS.search(txt):
                    continue
                nums = _NUM.findall(txt)
                if not nums:
                    continue
                rows.append({'段落号': f'tbl{ti}.r{ri}.c{ci}', '类型': 'cell', '内容': txt[:200], '数字数量': len(nums)})
                n += 1

    return rows, n


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    ap.add_argument('--out', default='01_extract_numbers.md')
    ap.add_argument('--json', default=None)
    args = ap.parse_args()

    rows, n = extract(args.docx)
    lines = [
        '# 结论性语句清单（extract_numbers.py v1）',
        f'来源：{args.docx} ｜ 提取条数：{n}',
        '',
        '| # | 类型 | 位置 | 数字数量 | 内容 |',
        '| --- | --- | --- | --- | --- |',
    ]
    for i, r in enumerate(rows, 1):
        lines.append(f'| {i} | {r["类型"]} | {r["段落号"]} | {r["数字数量"]} | {r["内容"]} |')
    Path(args.out).write_text('\n'.join(lines), encoding='utf-8')
    print(f'✅ 结论性语句清单已写出：{args.out}（{n} 条）')
    if args.json:
        Path(args.json).write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding='utf-8')


if __name__ == '__main__':
    main()