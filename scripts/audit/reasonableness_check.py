# -*- coding: utf-8 -*-
"""
reasonableness_check.py —— ⑥ 合理性门禁（v1）

环比 ±50% 与新增/移除标的必须在变更摘要.md 逐条点名。
依赖：python-docx
"""
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError as e:
    raise SystemExit('❌ python-docx 未安装') from e

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

_NUM = re.compile(r'-?\d+(?:\.\d+)?')


def _parse_num(s):
    m = _NUM.search(str(s))
    if not m:
        return None
    try:
        return float(m.group(0))
    except ValueError:
        return None


def _row_data(table, ki):
    out = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if ki >= len(cells):
            continue
        out.append(cells)
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('old_docx')
    ap.add_argument('new_docx')
    ap.add_argument('--key-col-name', default='公司简称')
    ap.add_argument('--jump-threshold', type=float, default='0.5')
    ap.add_argument('--roster-note', default=None, help='变更摘要.md 路径')
    ap.add_argument('--out', default='06_reasonableness.md')
    args = ap.parse_args()

    if not Path(args.new_docx).exists():
        Path(args.out).write_text(
            f'# 合理性报告（reasonableness_check.py v1）\n\n新月报未生成（{args.new_docx}）——门禁 SKIPPED。\n',
            encoding='utf-8')
        print(f'⚠️  合理性：SKIPPED（新月报不存在）')
        sys.exit(2)

    old = Document(args.old_docx)
    new = Document(args.new_docx)

    # 旧/新业务键集合
    def _keys(doc):
        ks = {}
        for t in doc.tables:
            if not t.rows:
                continue
            h = [c.text.strip() for c in t.rows[0].cells]
            if args.key_col_name not in h:
                continue
            ki = h.index(args.key_col_name)
            for r in t.rows[1:]:
                cells = [c.text.strip() for c in r.cells]
                if ki < len(cells):
                    ks[cells[ki]] = cells
        return ks

    a = _keys(old)
    b = _keys(new)
    added = [k for k in b if k not in a]
    removed = [k for k in a if k not in b]

    # 环比跳跃（仅按"旧/新同一行"做粗略对比）
    jumps = []
    for k in (set(a.keys()) & set(b.keys())):
        for i in range(1, min(len(a[k]), len(b[k]))):
            av = _parse_num(a[k][i])
            bv = _parse_num(b[k][i])
            if av is None or bv is None or av == 0:
                continue
            change = (bv - av) / abs(av)
            if abs(change) >= args.jump_threshold:
                jumps.append(f'{k} · 列{i}：{av} → {bv}（{change:+.1%}）')

    # 变更摘要是否点名
    summary = ''
    if args.roster_note and Path(args.roster_note).exists():
        summary = Path(args.roster_note).read_text(encoding='utf-8')

    hard = []
    for k in added + removed:
        if k not in summary:
            hard.append(f'{"新增" if k in added else "移除"}标的「{k}」未在变更摘要点名')
    for j in jumps:
        if not any(tok in summary for tok in j.split(' · ')[0:1]):
            hard.append(f'环比跳跃 {j} 未在变更摘要交代')

    lines = [
        '# 合理性校验报告（reasonableness_check.py v1）',
        f'环比阈值：±{args.jump_threshold:.0%} ｜ 变更摘要：{args.roster_note or "(未提供)"}',
        f'新增：{len(added)} ｜ 移除：{len(removed)} ｜ 环比跳跃：{len(jumps)}',
        '',
    ]
    if added:
        lines.append('## 新增标的')
        for k in added:
            lines.append(f'- {k}')
    if removed:
        lines.append('## 移除标的')
        for k in removed:
            lines.append(f'- {k}')
    if jumps:
        lines.append('## 环比跳跃')
        for j in jumps:
            lines.append(f'- {j}')
    if hard:
        lines += ['', f'## ❌ 硬伤：{len(hard)} 项']
        for h in hard:
            lines.append(f'- {h}')
    else:
        lines += ['', '## ✅ 无硬伤（新增/移除/跳跃均已在变更摘要交代）']

    Path(args.out).write_text('\n'.join(lines), encoding='utf-8')

    if hard:
        print(f'❌ 合理性：{len(hard)} 项硬伤（{args.out}）')
        sys.exit(1)
    print(f'✅ 合理性：无硬伤（{args.out}）')


if __name__ == '__main__':
    main()