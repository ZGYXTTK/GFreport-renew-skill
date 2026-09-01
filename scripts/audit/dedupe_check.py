# -*- coding: utf-8 -*-
"""
dedupe_check.py —— ⑪ 跨月去重门禁（v1，2026-08 期事故驱动，评审报告 🔴#3）

事故案例：
  - 经纬航宇：上期 7-4 天使轮 1.00 亿（闻名投资/硬核坚果），本期 8-1 又报一次
    「天使轮 1 亿」同额同投资方——上期事件重复计入；
  - 微光启航：上期 7-7 天使轮 1.00 亿（长风资本领投），本期 8-18 同轮报「0.3 亿 /
    北方某产业机构」——金额与投资方互相矛盾，且出现不可接受的无名主体。

本门禁做法：
  1. 在上期/新月报中启发式定位「投融资表」（表头含 轮次/融资/投资方/金额）；
  2. 每行提取指纹 = （归一主体名, 轮次）+ 金额（统一折算万元）；
  3. 跨月比对：
       同指纹 + 金额相对差 > tol      → ❌ 硬伤（金额矛盾）
       同指纹 + 金额一致（或均缺失）  → ⚠️ 警告（上期事件疑似重复计入，须在变更摘要说明）
       本期行含「某机构/某产业/匿名」 → ❌ 硬伤（无名主体不可入库）
  4. 找不到投融资表 → SKIPPED（exit 2，由盲审人工确认）。

退出码：0=无硬伤 / 1=存在硬伤 / 2=SKIPPED
"""
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError as e:
    raise SystemExit('❌ python-docx 未安装：pip install python-docx') from e

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

_NUM = re.compile(r'[-+]?\d[\d,]*\.?\d*')
_ROUND = re.compile(
    r'(天使|种子|Pre-?[-‐]?A\d?|A\+?|B\+?|C\+?|D\+?|E\+?|F\+?|战略投资|战略融资|股权融资|'
    r'定向增发|定增|可转债|增资|IPO|上市|并购|老股转让|新三板|挂牌|Pre-?[-‐]?IPO)',
    re.IGNORECASE)
_ANON = re.compile(r'某[机构产业公司基金资本方投资者投资]|[匿名未具]名|不具名|未披露名称|匿名主体')
_ORG_SUFFIX = ('股份有限公司', '有限责任公司', '有限公司', '控股集团', '科技集团', '（集团）', '(集团)')


def _norm_name(s):
    s = re.sub(r'[\s（）()【】\[\]·]', '', str(s or ''))
    for suf in _ORG_SUFFIX:
        if s.endswith(suf):
            s = s[:-len(suf)]
            break
    return s


def _amount_wan(text):
    """提取第一个 数字+亿/万 折算为万元；纯数字按万元计。返回 float 或 None。"""
    m = re.search(r'([-+]?\d[\d,]*\.?\d*)\s*(亿|万)?\s*元', str(text or ''))
    if not m:
        m2 = _NUM.search(str(text or ''))
        if not m2:
            return None
        return float(m2.group(0).replace(',', ''))
    v = float(m.group(1).replace(',', ''))
    if m.group(2) == '亿':
        return v * 10000
    return v


def _cell_texts(row):
    return [c.text.strip() for c in row.cells]


def _find_funding_tables(doc):
    """启发式定位全部一级市场投融资表（含跨页分表）。
    判据：表头（前两行）含 获投日期/投资方/轮次 之一，且含「公司」；
    排除二级市场再融资（融资方式）、IPO 在审（审核状态）、行情（涨跌幅）、辅导（辅导机构）表。
    返回 [(table, ti), ...]。"""
    groups = []
    exclude_kw = ('融资方式', '审核状态', '涨跌幅', '辅导机构', '上市板块')
    for ti, t in enumerate(doc.tables):
        if len(t.rows) < 3:
            continue
        head = ' '.join(' '.join(_cell_texts(r)) for r in t.rows[:2])
        if not any(k in head for k in ('获投日期', '投资方', '轮次')):
            continue
        if '公司' not in head:
            continue
        if any(k in head for k in exclude_kw):
            continue
        groups.append((t, ti))
    return groups


def _extract_rows(table):
    """提取 [(name_norm, name_raw, round, amount_wan, anon, row_idx), ...]
    金额只从表头「金额/融资额/募资」列取——避免误抓日期/序号数字（如 2026-08-01 → 2026）。"""
    out = []
    header = _cell_texts(table.rows[0])
    name_col = next((i for i, h in enumerate(header)
                     if any(k in h for k in ('公司', '主体', '企业', '项目', '标的', '名称'))), 0)
    amount_col = next((i for i, h in enumerate(header)
                       if i != name_col and any(k in h for k in ('金额', '融资额', '募资'))), None)
    for ri, row in enumerate(table.rows[1:], start=1):
        cells = _cell_texts(row)
        if not any(cells):
            continue
        raw_name = cells[name_col] if name_col < len(cells) else (cells[0] if cells else '')
        name_norm = _norm_name(raw_name)
        if not name_norm or _NUM.search(raw_name or ''):
            # 主体列误配数字列时回退：取行内第一个不含数字的 2-20 字文本
            alt = next((c for c in cells if c and not _NUM.search(c) and 2 <= len(c) <= 20), None)
            if not alt:
                continue
            raw_name, name_norm = alt, _norm_name(alt)
        rm = _ROUND.search(' '.join(cells))
        round_name = rm.group(0) if rm else ''
        amt = None
        if amount_col is not None and amount_col < len(cells):
            amt = _amount_wan(cells[amount_col])
        anon = bool(_ANON.search(' '.join(cells)))
        out.append((name_norm, raw_name, round_name, amt, anon, ri))
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--old', required=True, help='上期月报 docx')
    ap.add_argument('--new', required=True, help='新月报 docx')
    ap.add_argument('--amount-tol', type=float, default=0.2,
                    help='同指纹金额相对差容差（默认 0.2 = 20%%）')
    ap.add_argument('--out', default='11_dedupe.md')
    args = ap.parse_args()

    for label, p in (('上期', args.old), ('本期', args.new)):
        if not Path(p).exists():
            Path(args.out).write_text(
                f'# 跨月去重报告（dedupe_check.py v1）\n\n{label}月报不存在（{p}）——门禁 SKIPPED。\n',
                encoding='utf-8')
            print(f'⚠️  跨月去重：SKIPPED（{label}月报不存在）')
            sys.exit(2)

    old_doc, new_doc = Document(args.old), Document(args.new)
    old_ts = _find_funding_tables(old_doc)
    new_ts = _find_funding_tables(new_doc)
    if not old_ts or not new_ts:
        Path(args.out).write_text(
            '# 跨月去重报告（dedupe_check.py v1）\n\n'
            f'未定位到投融资表（old_tables={len(old_ts)}, new_tables={len(new_ts)}）——门禁 SKIPPED，'
            '请盲审人工核对跨月重复与无名主体。\n', encoding='utf-8')
        print('⚠️  跨月去重：SKIPPED（未定位到投融资表）')
        sys.exit(2)

    old_rows = [r for t, _ti in old_ts for r in _extract_rows(t)]
    new_rows = [r for t, _ti in new_ts for r in _extract_rows(t)]
    old_map = {}
    for r in old_rows:
        old_map.setdefault((r[0], r[2]), []).append(r)

    hard, warn = [], []
    checked = 0
    for name_norm, raw, round_name, amt, anon, ri in new_rows:
        if anon:
            hard.append(f'第{ri + 1}行「{raw}」含无名主体（某机构/匿名类表述）——不可入库，须具名或删行')
        if not round_name:
            continue
        key = (name_norm, round_name)
        old_matches = old_map.get(key, [])
        if not old_matches:
            # 宽松匹配：轮次相同 + 主体名相互包含（≥4 字，兼容「经纬航宇」vs「经纬航宇（上海）」）
            for (on, oround), rows in old_map.items():
                if oround != round_name:
                    continue
                if len(name_norm) >= 4 and len(on) >= 4 and (name_norm in on or on in name_norm):
                    old_matches = rows
                    break
        if old_matches:
            checked += 1
        for old_r in old_matches:
            old_amt = old_r[3]
            if amt is not None and old_amt is not None and old_amt != 0:
                diff = abs(amt - old_amt) / max(abs(amt), abs(old_amt))
                if diff > args.amount_tol:
                    hard.append(f'「{raw}」{round_name}：金额矛盾 上期 {old_amt:.0f} 万 vs 本期 {amt:.0f} 万'
                                f'（相对差 {diff:.0%} > {args.amount_tol:.0%}）——须以权威源更正并在变更摘要说明')
                else:
                    warn.append(f'「{raw}」{round_name} 金额与上期一致（{amt:.0f} 万）——上期事件疑似重复计入，'
                                f'须在变更摘要标注「系上期事件更正/延续」或删行')
            else:
                warn.append(f'「{raw}」{round_name} 与上期同主体同轮次（金额缺失可比对）——疑似重复计入，请人工确认')

    n_new = len(new_rows)
    lines = [
        '# 跨月去重报告（dedupe_check.py v1）',
        '',
        '| 指标 | 值 |',
        '| --- | --- |',
        f'| 上期投融资表 | {len(old_ts)} 张（分页合并），{len(old_rows)} 行 |',
        f'| 本期投融资表 | {len(new_ts)} 张（分页合并），{n_new} 行 |',
        f'| 同主体同轮次命中 | {checked} |',
        f'| ⚠️ 疑似重复 | {len(warn)} |',
        f'| ❌ 硬伤 | {len(hard)} |',
        '',
    ]
    if warn:
        lines += [f'## ⚠️ 疑似重复计入：{len(warn)} 项', '']
        lines += [f'- {w}' for w in warn]
    if hard:
        lines += ['', f'## ❌ 硬伤：{len(hard)} 项', '']
        lines += [f'- {h}' for h in hard]
    else:
        lines += ['', '## ✅ 无金额矛盾与无名主体']
    lines += ['', '### 方法说明',
              '- 指纹 =（归一主体名, 轮次）；金额统一折算万元；相对差超容差判矛盾；',
              '- 金额一致的同指纹判「疑似重复」——允许「上期事件更正/延续」，但必须在变更摘要显式说明。',
              ]

    Path(args.out).write_text('\n'.join(lines), encoding='utf-8')
    print(f'跨月去重：本期 {n_new} 行，同指纹命中 {checked}，⚠️{len(warn)} ❌{len(hard)}（{args.out}）')
    if hard:
        sys.exit(1)


if __name__ == '__main__':
    main()
