# -*- coding: utf-8 -*-
"""
diff_empty.py —— ③ 空值 diff 门禁（v1）

业务键对齐：旧月报有值 → 新月报无值 = 硬伤。
依赖：python-docx
"""
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError as e:
    raise SystemExit('❌ python-docx 未安装') from e

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')


def _norm_key(s):
    s = re.sub(r'[（(][^）)]*[）)]', '', s or '')
    s = re.sub(r'(股份有限公司|有限责任公司|有限公司|控股集团|集团)', '', s)
    return s.strip()


def _read_table_keyed(doc, key_col_name):
    """返回 {key_norm: [cell_text, ...]}"""
    out = {}
    for table in doc.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if key_col_name not in header:
            continue
        ki = header.index(key_col_name)
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if ki >= len(cells):
                continue
            k = _norm_key(cells[ki])
            if not k:
                continue
            out[k] = cells
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('old_docx')
    ap.add_argument('new_docx')
    ap.add_argument('--key-col-name', default='公司简称')
    ap.add_argument('--out', default='03_diff_empty.md')
    args = ap.parse_args()

    if not Path(args.new_docx).exists():
        print(f'⚠️  新月报不存在：{args.new_docx}（门禁 SKIPPED）')
        Path(args.out).write_text(
            f'# 空值 diff 报告（diff_empty.py v1）\n\n新月报尚未生成（{args.new_docx}）——门禁 SKIPPED。\n',
            encoding='utf-8')
        sys.exit(2)  # SKIPPED

    old = Document(args.old_docx)
    new = Document(args.new_docx)
    a = _read_table_keyed(old, args.key_col_name)
    b = _read_table_keyed(new, args.key_col_name)

    hard = []
    warn = []
    keys = set(a.keys()) | set(b.keys())
    for k in keys:
        a_cells = a.get(k)
        b_cells = b.get(k)
        if not a_cells:
            continue
        if not b_cells:
            hard.append(f'新⽉报缺失键：{k}')
            continue
        # 旧有值 → 新空值
        for i in range(min(len(a_cells), len(b_cells))):
            if a_cells[i].strip() and not b_cells[i].strip():
                hard.append(f'{k} · 第{i}列：旧「{a_cells[i]}」 → 新空值')

    lines = [
        '# 空值 diff 报告（diff_empty.py v1）',
        f'旧：{args.old_docx} ｜ 新：{args.new_docx} ｜ 业务键：{args.key_col_name}',
        f'旧表键数：{len(a)} ｜ 新表键数：{len(b)} ｜ 交集：{len(set(a.keys()) & set(b.keys()))}',
        '',
    ]
    if hard:
        lines.append(f'## ❌ 硬伤：{len(hard)} 项')
        for h in hard:
            lines.append(f'- {h}')
    else:
        lines.append('## ✅ 无硬伤（旧有值 → 新空值 = 0）')

    Path(args.out).write_text('\n'.join(lines), encoding='utf-8')

    if hard:
        print(f'❌ 空值 diff：{len(hard)} 项硬伤（{args.out}）')
        sys.exit(1)
    print(f'✅ 空值 diff：无硬伤（{args.out}）')


if __name__ == '__main__':
    main()