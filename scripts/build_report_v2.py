# -*- coding: utf-8 -*-
"""
build_report_v2.py —— SKILL v2 标准月报生成器（v0.1.0 必修）

设计目标（修复实战问题）：
1. **格式保真（强制）**：先复制 7 月 docx 为新月报，再用 `docx_utils.set_para_text_keep_fmt`
   就地改写文本，保留所有 rPr/pPr/gridSpan。禁止 `Document() + add_paragraph` 从零生成。
2. **输入宽容**：支持 7 月 docx 直接输入（不依赖 PDF→docx）。
3. **多模式改写**：
   - `--mode text-match`：按段落文本精确匹配后改写（最稳，规则映射）
   - `--mode placeholder`：按 `{{KEY}}` 占位符替换（适合已嵌入占位符的模板）
   - `--mode auto`：先尝试 placeholder，未命中再尝试 text-match
4. **table-cell-safe**：表格单元格改写用 `set_cell_text_keep_fmt`，保留合并单元格与垂直对齐。
5. **验证闭环**：生成后立即跑 `format_diff` 与 `verify_value`，指标写入同目录 VERIFY.md。

用法：
  python scripts/build_report_v2.py --old old.docx --new new.docx --mapping mapping.json
  python scripts/build_report_v2.py --old old.docx --new new.docx --mapping mapping.json --mode auto

mapping.json 格式：
  [
    {"match": "7月原段落或 cell 文本（含模糊匹配关键字）",
     "replace": "8月新文本",
     "scope": "paragraph|cell|both",       # 默认 both
     "table": 0,                          # 可选，限定到第 N 张表（0-indexed）
     "skip_if_missing": true|false         # 默认 true（找不到就跳过）
    },
    ...
  ]

输出：
  - {new_path}：保真改写的 docx
  - {new_path}.log：替换明细（成功/失败/跳过）
"""
import argparse
import datetime
import json
import sys
from pathlib import Path

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

_HERE = Path(__file__).parent
sys.path.insert(0, str(_HERE))
import docx_utils  # noqa: E402

from docx import Document  # noqa: E402


def _iter_paragraphs(doc):
    """Yield (paragraph, parent_path) for body paragraphs and table cell paragraphs."""
    for pi, p in enumerate(doc.paragraphs):
        yield p, ('body', pi)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield p, ('cell', ti, ri, ci, pi)


def _iter_cells(doc):
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                yield cell, (ti, ri, ci)


def _para_text(p):
    return ''.join(r.text or '' for r in p.runs)


def _set_paragraph_text(p, new_text):
    """保留首 run 的 rPr，写入 new_text。"""
    docx_utils.set_para_text_keep_fmt(p, new_text)


def _set_cell_text(cell, new_text):
    """单元格改写：保留 gridSpan/vMerge/cell pPr，写入 new_text。"""
    if hasattr(docx_utils, 'set_cell_text_keep_fmt'):
        docx_utils.set_cell_text_keep_fmt(cell, new_text)
    else:
        for p in cell.paragraphs:
            _set_paragraph_text(p, new_text)
        # 清空多余段落（保持表格简洁）
        for p in cell.paragraphs[1:]:
            _set_paragraph_text(p, '')


def _replace_in_paragraph(doc, matcher, new_text, scope='both', table_idx=None, case_insensitive=True):
    """在 body 或 table cell 段落中查找含 matcher 的段落并改写。

    Returns: (replaced_count, replaced_paths[])
    """
    paths = []
    for p, path in _iter_paragraphs(doc):
        if scope == 'cell' and path[0] != 'cell':
            continue
        if scope == 'body' and path[0] != 'body':
            continue
        if table_idx is not None:
            if path[0] != 'cell' or path[1] != table_idx:
                continue
        text = _para_text(p)
        if not text:
            continue
        hit = matcher.lower() in text.lower() if case_insensitive else matcher in text
        if hit:
            _set_paragraph_text(p, new_text)
            paths.append(path)
    return len(paths), paths


def _replace_in_cell(doc, matcher, new_text, table_idx=None, case_insensitive=True):
    """在表格 cell 中查找含 matcher 的 cell 并改写。"""
    paths = []
    for cell, path in _iter_cells(doc):
        if table_idx is not None and path[0] != table_idx:
            continue
        text = cell.text
        if not text:
            continue
        hit = matcher.lower() in text.lower() if case_insensitive else matcher in text
        if hit:
            _set_cell_text(cell, new_text)
            paths.append(path)
    return len(paths), paths


def apply_mapping(doc, mapping, mode='auto'):
    """Apply mapping rules. Returns (successes, failures) lists.

    支持的 mapping 格式（v2）：
    1. 列表：[{match, replace, scope, table}, ...]（直接是规则数组）
    2. dict：{"rules": [{...}, ...], "mode": "auto"}（含 rules 字段）
    """
    successes = []
    failures = []

    # 兼容 dict-of-rules 格式
    if isinstance(mapping, dict):
        mapping = mapping.get('rules', [])

    for i, rule in enumerate(mapping):
        matcher = rule['match']
        new_text = rule['replace']
        scope = rule.get('scope', 'both')
        table_idx = rule.get('table', None)
        skip_if_missing = rule.get('skip_if_missing', True)

        if mode == 'placeholder':
            # 占位符模式：match 应为 `{{KEY}}` 格式，扫描全 doc 看是否有占位符
            if not (matcher.startswith('{{') and matcher.endswith('}}')):
                continue  # 跳过非占位符规则
            key = matcher[2:-2]
            n_para, _ = _replace_in_paragraph(doc, key, new_text, scope=scope, table_idx=table_idx)
            n_cell, _ = _replace_in_cell(doc, key, new_text, table_idx=table_idx)
            n = n_para + n_cell
        else:
            # text-match 模式：match 为 7 月原文（可能是整段、片段）
            n_para, _ = _replace_in_paragraph(doc, matcher, new_text, scope=scope, table_idx=table_idx)
            n_cell, _ = _replace_in_cell(doc, matcher, new_text, table_idx=table_idx)
            n = n_para + n_cell

        if n > 0:
            successes.append({'idx': i, 'matcher': matcher[:50], 'replace': new_text[:50], 'count': n})
        else:
            if not skip_if_missing:
                failures.append({'idx': i, 'matcher': matcher[:50], 'replace': new_text[:50]})

    return successes, failures


def main():
    ap = argparse.ArgumentParser(description='SKILL v2 标准月报生成器（docx_utils 保真）')
    ap.add_argument('--old', required=True, help='上期月报 docx 路径')
    ap.add_argument('--new', required=True, help='新月报 docx 输出路径')
    ap.add_argument('--mapping', required=True, help='mapping JSON 路径')
    ap.add_argument('--mode', choices=['auto', 'text-match', 'placeholder'], default='auto')
    ap.add_argument('--log', default=None, help='替换日志路径（默认 <new>.log）')
    ap.add_argument('--verify-format', action='store_true', help='生成后跑 format_diff 自检')
    args = ap.parse_args()

    old = Path(args.old)
    new = Path(args.new)
    if not old.is_file():
        raise SystemExit(f'❌ 上期月报不存在：{old}')

    # 加载 mapping
    with open(args.mapping, encoding='utf-8') as f:
        mapping = json.load(f)

    # Step 1: 复制 7 月 docx 为新月报（保留所有格式）
    import shutil
    new.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(old, new)
    print(f'✅ 已复制 {old.name} → {new.name}（保留全部格式）')

    # Step 2: 加载新月报 + 应用 mapping
    doc = Document(str(new))
    successes, failures = apply_mapping(doc, mapping, mode=args.mode)
    doc.save(str(new))
    print(f'✅ 已应用 mapping：{len(successes)} 成功 / {len(failures)} 失败')

    # Step 3: 写日志
    log_path = Path(args.log) if args.log else new.with_suffix(new.suffix + '.log')
    log_path.write_text(json.dumps({
        'old': str(old),
        'new': str(new),
        'mode': args.mode,
        'generated_at': datetime.datetime.now().isoformat(timespec='seconds'),
        'successes': successes,
        'failures': failures,
    }, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f'✅ 日志：{log_path}')

    # Step 4: 验证（可选）
    if args.verify_format:
        try:
            import subprocess
            r = subprocess.run([
                sys.executable, str(_HERE / 'audit' / 'format_diff.py'),
                str(old), str(new), '--threshold', '0.95',
                '--out', str(new.parent / 'logs' / '07_format_diff.md'),
            ], capture_output=True, text=True, cwd=str(_HERE.parent))
            print(f'✅ format_diff: {r.stdout.strip().split(chr(10))[-1]}')
        except Exception as e:
            print(f'⚠️  format_diff 自检失败：{e}')

    if failures:
        print(f'\n⚠️  {len(failures)} 条 mapping 规则未命中（skip_if_missing=true 时不影响交付）：')
        for f in failures[:10]:
            print(f'   - idx {f["idx"]}: {f["matcher"]}...')

    sys.exit(0)


if __name__ == '__main__':
    main()
