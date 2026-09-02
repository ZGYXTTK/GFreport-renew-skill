# -*- coding: utf-8 -*-
"""
docx_utils.py —— 格式保真工具库（v1）

保证 docx 改写后：
  1. 多 run 单元格：逐 run 克隆 rPr（避免「整格变粗/变字号」）
  2. 多段落单元格：逐段克隆 pPr
  3. 合并单元格：add_row 时按网格重建 gridSpan / vMerge
  4. 超链接 run（<w:hyperlink>）内嵌的 run 改写不留旧文字

依赖：python-docx（本机已装；缺失时 install 报错而非静默）。
"""
from copy import deepcopy

try:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError as e:
    raise SystemExit('❌ python-docx 未安装：pip install python-docx') from e


# ---------- 基础 XML 工具 ----------
def _find(el, tag):
    return el.find(qn(tag))


def _remove(el, tag):
    old = _find(el, tag)
    if old is not None:
        el.remove(old)
    return old


def _clone_into(el, tag, src_el):
    """把 src_el 的指定子元素深拷贝插入 el（替换同名旧子元素）。"""
    src = _find(src_el, tag)
    if src is None:
        return
    _remove(el, tag)
    el.insert(0, deepcopy(src))


# ---------- 段落 ----------
def _all_run_elements(p):
    """返回段落下所有 <w:r> 元素，含 <w:hyperlink> 内部的 run，按文档顺序。"""
    els = []
    for child in p._element:
        if child.tag == qn('w:r'):
            els.append(child)
        elif child.tag == qn('w:hyperlink'):
            for sub in child:
                if sub.tag == qn('w:r'):
                    els.append(sub)
    return els


def _set_run_text(el, text):
    """向 <w:r> 写入文本（复用首个 w:t，清空其余）。"""
    ts = el.findall(qn('w:t'))
    if not ts:
        t = el.makeelement(qn('w:t'), {})
        t.text = text
        el.append(t)
    else:
        ts[0].text = text
        for t in ts[1:]:
            t.text = ''


def set_para_text_keep_fmt(p, new_text):
    """改段落文字但保留段落格式(pPr)与首个 run 格式(rPr)，兼容超链接内 run。"""
    runs = _all_run_elements(p)
    if runs:
        _set_run_text(runs[0], new_text)
        for el in runs[1:]:
            _set_run_text(el, '')
    else:
        p.add_run(new_text)


def _apply_rpr_overrides(run, overrides):
    """把 overrides（如 {'b': True} 或 {'sz': 21}）应用到 run 的 rPr。"""
    rpr = _find(run._element, 'w:rPr')
    if rpr is None:
        rpr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rpr)
    for k, v in overrides.items():
        tag = qn('w:%s' % k)
        el = rpr.find(tag)
        if v is True:
            if el is None:
                el = rpr.makeelement(tag, {})
                rpr.append(el)
        elif v is False or v is None:
            if el is not None:
                rpr.remove(el)
        else:
            if el is None:
                el = rpr.makeelement(tag, {})
                rpr.append(el)
            el.set(qn('w:val'), str(v))


def set_para_segments_keep_fmt(p, segments):
    """
    按多段格式回填段落（「标题加粗 run + 正文不加粗 run」这类）。
    segments: list[tuple[str, dict|None]]，dict 为可选 rPr 覆盖（如 {"b": True}）。
    """
    if not p.runs:
        p.add_run('')
    base_rpr = deepcopy(_find(p.runs[0]._element, 'w:rPr'))
    for el in _all_run_elements(p):
        el.getparent().remove(el)
    for text, override in segments:
        run = p.add_run(text)
        if base_rpr is not None:
            run._element.insert(0, deepcopy(base_rpr))
        if override:
            _apply_rpr_overrides(run, override)


# ---------- 单元格 ----------
def set_cell_text_keep_fmt(cell, value):
    """改单元格文字但保留 tcPr(vAlign)/pPr(jc)/rPr(字号)，兼容超链接内 run。"""
    runs = _all_run_elements(cell.paragraphs[0])
    if runs:
        _set_run_text(runs[0], value)
        for el in runs[1:]:
            _set_run_text(el, '')
    else:
        cell.paragraphs[0].add_run(value)


def _clone_paragraphs(src_cell, dst_cell):
    for p in list(dst_cell.paragraphs):
        p._element.getparent().remove(p._element)
    for sp in src_cell.paragraphs:
        new_p = dst_cell.add_paragraph()
        _clone_into(new_p._element, 'w:pPr', sp._element)
        for r in sp.runs:
            nr = new_p.add_run(r.text)
            _clone_into(nr._element, 'w:rPr', r._element)


def copy_cell_fmt(src_cell, dst_cell):
    """深拷贝单元格全套格式：tcPr（含 vAlign、gridSpan、vMerge）+ 每段 pPr + 每 run rPr。"""
    _clone_into(dst_cell._tc, 'w:tcPr', src_cell._tc)
    _clone_paragraphs(src_cell, dst_cell)


def _row_grid(row):
    """返回该行按网格展开的 [(tc_element, grid_index, span), ...]。"""
    result = []
    grid_idx = 0
    for tc in row._tr.tc_lst:
        tcPr = _find(tc, 'w:tcPr')
        span = 1
        if tcPr is not None:
            gs = _find(tcPr, 'w:gridSpan')
            if gs is not None:
                span = int(gs.get(qn('w:val'), '1'))
        result.append((tc, grid_idx, span))
        grid_idx += span
    return result


def add_row_copy_fmt(table, values, src_row_index=None):
    """
    新增数据行：按源行「逻辑列」重建单元格（含 gridSpan），再复制格式填值。
    values 按源行「逻辑列」顺序；src_row_index 不传时取最后一条数据行作模板。
    """
    if len(table.rows) < 2:
        raise ValueError('表格需至少包含表头 + 一条数据行作为格式模板')
    src_row = table.rows[src_row_index if src_row_index is not None else len(table.rows) - 1]
    src_grid = _row_grid(src_row)
    if len(values) != len(src_grid):
        raise ValueError('values 数量 %d 与源行逻辑列数 %d 不一致' % (len(values), len(src_grid)))

    tr = table.add_row()
    for tc in list(tr._tr.tc_lst):
        tr._tr.remove(tc)
    for v, (src_tc, _grid_idx, _span) in zip(values, src_grid):
        tc = OxmlElement('w:tc')
        tr._tr.append(tc)
        _clone_into(tc, 'w:tcPr', src_tc)
        _fill_tc_from_template(tc, src_tc, v)
    return tr


def _fill_tc_from_template(dst_tc, src_tc, value):
    """以源单元格第一段为模板：复制 pPr + 首 run rPr，回填文字到新建 run。"""
    p = OxmlElement('w:p')
    dst_tc.append(p)
    src_ps = src_tc.findall(qn('w:p'))
    if src_ps:
        _clone_into(p, 'w:pPr', src_ps[0])
    src_rs = src_ps[0].findall(qn('w:r')) if src_ps else []
    r = OxmlElement('w:r')
    p.append(r)
    if src_rs:
        _clone_into(r, 'w:rPr', src_rs[0])
    t = OxmlElement('w:t')
    t.text = str(value)
    r.append(t)


if __name__ == '__main__':
    print('docx_utils loaded. Imports OK.')